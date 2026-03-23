from app.models.report_request import ReportRequest
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

G2T_BLUE = RGBColor(0x2E, 0x5F, 0xA3)

# Severity colours for Mängel table rows
SCHWERE_COLORS = {
    "minor":    "FFF3CD",  # amber-light
    "major":    "FFD8A8",  # orange-light
    "critical": "F8D7DA",  # red-light
}

SCHWERE_LABELS = {
    "minor":    {"de": "Gering",    "en": "Minor",    "es": "Leve"},
    "major":    {"de": "Erheblich", "en": "Major",    "es": "Grave"},
    "critical": {"de": "Kritisch",  "en": "Critical", "es": "Crítico"},
}

NO_WORKS_REASON_LABELS = {
    "kein_arbeitstag":     {"de": "Kein geplanter Arbeitstag",        "en": "No working day planned",          "es": "Día no laborable"},
    "havarie":             {"de": "Havarie / Betriebsstörung",         "en": "Breakdown / operational incident", "es": "Avería / incidente operativo"},
    "extremwetter":        {"de": "Extremwetter / Unwetter",           "en": "Extreme weather / storm",          "es": "Clima extremo / tormenta"},
    "sicherheit":          {"de": "Terrorlage / Sicherheitsereignis",  "en": "Security event",                   "es": "Evento de seguridad"},
    "technischer_ausfall": {"de": "Technischer Ausfall",               "en": "Technical failure",                "es": "Fallo técnico"},
    "sonstiges":           {"de": "Sonstiges",                         "en": "Other",                            "es": "Otros"},
}


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)


def set_title_bg(title, hex_color: str):
    pPr = title._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)


def add_section_heading(doc, text: str):
    h = doc.add_heading(text, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = G2T_BLUE
    return h


def build_docx(request: ReportRequest, report_text: str) -> bytes:
    doc = Document()
    lang_key = request.report_language if request.report_language in ("de", "en", "es") else "de"

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ──────────────────────────────────────────────────────────────────
    title_text = {"de": "BAUTAGESBERICHT", "en": "DAILY CONSTRUCTION REPORT", "es": "INFORME DE OBRA DIARIA"}
    title = doc.add_heading(title_text.get(lang_key, "BAUTAGESBERICHT"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        set_title_bg(title, '2E5FA3')
        title.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        title.runs[0].font.size = Pt(20)

    # ── No-works banner (if applicable) ───────────────────────────────────────
    if request.no_works and request.no_works_reason and request.no_works_reason != "kein_arbeitstag":
        reason_label = NO_WORKS_REASON_LABELS.get(request.no_works_reason, {}).get(lang_key, request.no_works_reason)
        if request.no_works_reason_text:
            reason_label += f": {request.no_works_reason_text}"
        banner = doc.add_paragraph()
        banner_run = banner.add_run(f"⚠ KEIN ARBEITSTAG — {reason_label.upper()}")
        banner_run.font.bold = True
        banner_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        banner_run.font.size = Pt(11)
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Header info table ──────────────────────────────────────────────────────
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    weather_str = request.weather or "—"
    if request.temp_celsius is not None:
        weather_str += f", {request.temp_celsius}°C"

    label_map = {
        "de": ["Projekt", "Datum", "Arbeitszeit", "Bauüberwacher", "Wetter"],
        "en": ["Project", "Date", "Working Hours", "Site Supervisor", "Weather"],
        "es": ["Proyecto", "Fecha", "Horario", "Supervisor de Obra", "Clima"],
    }
    labels = label_map.get(lang_key, label_map["de"])
    info_data = [
        (labels[0], f"{request.project_name} ({request.project_id})"),
        (labels[1], request.date),
        (labels[2], f"{request.start_time or '--:--'} – {request.end_time or '--:--'}"),
        (labels[3], request.supervisor),
        (labels[4], weather_str),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i].cells
        row[0].text = label
        row[1].text = value
        set_cell_bg(row[0], "D6E4F7")
        for cell in row:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.bold = (cell == row[0])

    doc.add_paragraph("")

    # ── Report body (Claude-generated) ────────────────────────────────────────
    heading_map = {"de": "Durchgeführte Arbeiten", "en": "Work Performed", "es": "Trabajos Realizados"}
    add_section_heading(doc, heading_map.get(lang_key, "Durchgeführte Arbeiten"))
    lines = report_text.strip().split("\n\n")
    for para in lines:
        clean = para.strip()
        if clean.upper() in ("BAUTAGESBERICHT", "DAILY CONSTRUCTION REPORT", "INFORME DE OBRA DIARIA"):
            continue
        if clean:
            p = doc.add_paragraph(clean)
            p.paragraph_format.space_after = Pt(6)

    # ── Workforce + Equipment — skip entirely for kein_arbeitstag ─────────────
    if not (request.no_works and request.no_works_reason == "kein_arbeitstag"):

        wf_heading = {"de": "Arbeitskräfte", "en": "Workforce", "es": "Personal"}
        add_section_heading(doc, wf_heading.get(lang_key, "Arbeitskräfte"))
        if request.workforce:
            wf_table = doc.add_table(rows=1, cols=2)
            wf_table.style = "Table Grid"
            hdr = wf_table.rows[0].cells
            wf_col_labels = {"de": ["Funktion", "Anzahl"], "en": ["Role", "Count"], "es": ["Función", "Cantidad"]}
            for i, lbl in enumerate(wf_col_labels.get(lang_key, ["Funktion", "Anzahl"])):
                hdr[i].text = lbl
                set_cell_bg(hdr[i], "2E5FA3")
                hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                hdr[i].paragraphs[0].runs[0].font.bold = True
            for w in request.workforce:
                row = wf_table.add_row().cells
                row[0].text = w.role
                row[1].text = str(w.count)
        else:
            doc.add_paragraph("Keine Angaben" if lang_key == "de" else ("No data" if lang_key == "en" else "Sin datos"))

        eq_heading = {"de": "Geräte und Maschinen", "en": "Equipment & Machinery", "es": "Equipos y Maquinaria"}
        add_section_heading(doc, eq_heading.get(lang_key, "Geräte und Maschinen"))
        if request.equipment:
            eq_table = doc.add_table(rows=1, cols=2)
            eq_table.style = "Table Grid"
            hdr = eq_table.rows[0].cells
            eq_col_labels = {"de": ["Gerät / Maschine", "Anzahl"], "en": ["Equipment / Machine", "Count"], "es": ["Equipo / Máquina", "Cantidad"]}
            for i, lbl in enumerate(eq_col_labels.get(lang_key, ["Gerät / Maschine", "Anzahl"])):
                hdr[i].text = lbl
                set_cell_bg(hdr[i], "2E5FA3")
                hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                hdr[i].paragraphs[0].runs[0].font.bold = True
            for e in request.equipment:
                row = eq_table.add_row().cells
                row[0].text = e.name
                row[1].text = str(e.count)
        else:
            doc.add_paragraph("Keine Angaben" if lang_key == "de" else ("No data" if lang_key == "en" else "Sin datos"))

    # ── Abnahmen — only if filled ──────────────────────────────────────────────
    ab = request.abnahme
    if ab and ab.item:
        ab_heading = {"de": "Abnahmen", "en": "Inspections", "es": "Inspecciones"}
        add_section_heading(doc, ab_heading.get(lang_key, "Abnahmen"))
        result_label = {"pass": "Bestanden", "fail": "Nicht bestanden",
                        "conditional": "Bedingt bestanden"}.get(ab.result or "", ab.result or "—")
        ab_field_labels = {
            "de": ["Abnahmeobjekt", "Abgenommen von", "Uhrzeit", "Ergebnis", "Anmerkungen"],
            "en": ["Inspection Item", "Approved By", "Time", "Result", "Notes"],
            "es": ["Objeto de Inspección", "Aprobado por", "Hora", "Resultado", "Notas"],
        }
        fl = ab_field_labels.get(lang_key, ab_field_labels["de"])
        ab_rows = [
            (fl[0], ab.item or "—"),
            (fl[1], ab.approver or "—"),
            (fl[2], ab.time or "—"),
            (fl[3], result_label),
        ]
        if ab.notes:
            ab_rows.append((fl[4], ab.notes))
        ab_table = doc.add_table(rows=len(ab_rows), cols=2)
        ab_table.style = "Table Grid"
        for i, (label, value) in enumerate(ab_rows):
            row = ab_table.rows[i].cells
            row[0].text = label
            row[1].text = value
            set_cell_bg(row[0], "D6E4F7")
            if row[0].paragraphs[0].runs:
                row[0].paragraphs[0].runs[0].font.bold = True

    # ── Mängel — only if filled ───────────────────────────────────────────────
    if request.maengel:
        maengel_heading = {"de": "Festgestellte Mängel", "en": "Identified Defects", "es": "Defectos Identificados"}
        add_section_heading(doc, maengel_heading.get(lang_key, "Festgestellte Mängel"))

        mn_col_labels = {
            "de": ["Nr.", "Beschreibung", "Ort / km", "Verantwortlich", "Frist", "Schwere"],
            "en": ["No.", "Description", "Location / km", "Responsible Party", "Deadline", "Severity"],
            "es": ["Nº", "Descripción", "Ubicación / km", "Responsable", "Plazo", "Gravedad"],
        }
        cols = mn_col_labels.get(lang_key, mn_col_labels["de"])

        mn_table = doc.add_table(rows=1, cols=6)
        mn_table.style = "Table Grid"
        hdr = mn_table.rows[0].cells
        for i, lbl in enumerate(cols):
            hdr[i].text = lbl
            set_cell_bg(hdr[i], "2E5FA3")
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            hdr[i].paragraphs[0].runs[0].font.bold = True

        for i, m in enumerate(request.maengel, 1):
            schwere_label = SCHWERE_LABELS.get(m.schwere or "", {}).get(lang_key, m.schwere or "—")
            row = mn_table.add_row().cells
            row[0].text = str(i)
            row[1].text = m.beschreibung or "—"
            row[2].text = m.ort or "—"
            row[3].text = m.verantwortlich or "—"
            row[4].text = m.frist or "—"
            row[5].text = schwere_label
            # Colour-code severity column
            bg = SCHWERE_COLORS.get(m.schwere or "", "FFFFFF")
            set_cell_bg(row[5], bg)
            if row[5].paragraphs[0].runs:
                row[5].paragraphs[0].runs[0].font.bold = True

    # ── Störungen / Probleme — only if filled ─────────────────────────────────
    if request.stoerungen:
        st_heading = {"de": "Störungen / Probleme", "en": "Issues / Delays", "es": "Problemas / Retrasos"}
        add_section_heading(doc, st_heading.get(lang_key, "Störungen / Probleme"))
        doc.add_paragraph(request.stoerungen)

    # ── Besondere Vorkommnisse — only if filled ────────────────────────────────
    if request.besonderheiten:
        bv_heading = {"de": "Besondere Vorkommnisse", "en": "Special Incidents", "es": "Incidentes Especiales"}
        add_section_heading(doc, bv_heading.get(lang_key, "Besondere Vorkommnisse"))
        if request.bg_meldung:
            bg_para = doc.add_paragraph()
            bg_run = bg_para.add_run("⚠ BG-MELDUNG ERFORDERLICH — Arbeitsunfall gemeldet.")
            bg_run.font.bold = True
            bg_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            bg_run.font.size = Pt(11)
            doc.add_paragraph("")
        doc.add_paragraph(request.besonderheiten)

    # ── Nächste Schritte — only if filled ─────────────────────────────────────
    if request.next_steps:
        ns_heading = {"de": "Geplante Arbeiten (Folgetag)", "en": "Planned Work (Next Day)", "es": "Trabajos Planificados (Día Siguiente)"}
        add_section_heading(doc, ns_heading.get(lang_key, "Geplante Arbeiten (Folgetag)"))
        doc.add_paragraph(request.next_steps)

    # ── Signature ──────────────────────────────────────────────────────────────
    sig_label = {"de": "Unterschrift Bauüberwacher", "en": "Signature Site Supervisor", "es": "Firma Supervisor de Obra"}
    doc.add_paragraph("")
    doc.add_paragraph(
        f"{sig_label.get(lang_key, 'Unterschrift Bauüberwacher')}: ___________________________"
        f"          Datum: {request.date}"
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()