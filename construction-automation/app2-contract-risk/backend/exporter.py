"""
exporter.py — DOCX generation for App 2 export endpoints.

Mode A: export_risk_report_docx(data) → bytes
  Generates a formatted risk analysis report.

Mode B: export_stellungnahme_docx(data) → bytes
  Generates a formal Stellungnahme ready for the AG to sign.

Design: uses python-docx (same library as App 4). No external template file
needed — document structure is built programmatically. This keeps deployment
simple (no template assets to manage on Render).
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Colour constants (risk level) ─────────────────────────────────────────────
_RED    = RGBColor(0xDC, 0x26, 0x26)  # Tailwind red-600
_AMBER  = RGBColor(0xD9, 0x77, 0x06)  # Tailwind amber-600
_GREEN  = RGBColor(0x16, 0xA3, 0x4A)  # Tailwind green-600
_DARK   = RGBColor(0x1F, 0x29, 0x37)  # near-black for body

_RISK_COLOUR = {"HIGH": _RED, "MEDIUM": _AMBER, "LOW": _GREEN}
_ASSESS_COLOUR = {"REJECT": _RED, "NEGOTIATE": _AMBER, "ACCEPT": _GREEN}


def _set_font(run, size_pt: int = 11, bold: bool = False, colour: RGBColor = None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if colour:
        run.font.color.rgb = colour


def _add_meta_line(doc, label: str, value: str, colour: RGBColor = None):
    """Add a key–value line (bold label, normal value)."""
    p = doc.add_paragraph()
    r_label = p.add_run(f"{label}: ")
    _set_font(r_label, bold=True)
    r_value = p.add_run(value)
    _set_font(r_value, colour=colour)


# ── Mode A — Risk Report ──────────────────────────────────────────────────────

def export_risk_report_docx(data: dict) -> bytes:
    """
    Export Mode A analysis result as a formatted DOCX.

    data keys: clauses (list), summary (dict)
    """
    doc = Document()
    _set_page_margins(doc)

    # Title
    title = doc.add_heading("VOB/B Vertragsrisikoanalyse", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph(f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    # ── Summary block ──
    summary = data.get("summary", {})
    doc.add_heading("Zusammenfassung", level=2)

    risk_level = summary.get("overall_risk_level", "N/A")
    risk_score = summary.get("overall_risk_score", 0)
    colour = _RISK_COLOUR.get(risk_level, _DARK)

    _add_meta_line(doc, "Gesamtrisiko", f"{risk_level} ({risk_score}/100)", colour)
    _add_meta_line(doc, "Hohe Risikoklauseln", str(summary.get("high_risk_count", 0)))
    _add_meta_line(doc, "Mittlere Risikoklauseln", str(summary.get("medium_risk_count", 0)))
    doc.add_paragraph(summary.get("summary_text", ""))
    doc.add_paragraph()

    # ── Top-3 risky clauses ──
    top3 = summary.get("top_3_risky_clauses", [])
    if top3:
        doc.add_heading("Top-Risikoklauseln", level=3)
        for i, c in enumerate(top3, 1):
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(f"{c['number']} — {c['title']}")
            r.bold = True
            r.font.color.rgb = _RED
            p.add_run(f": {c.get('reason', '')}")

    doc.add_page_break()

    # ── Per-clause details ──
    doc.add_heading("Klauseldetails", level=2)

    for clause in data.get("clauses", []):
        risk = clause.get("risk_level", "low").upper()
        colour = _RISK_COLOUR.get(risk, _DARK)

        h = doc.add_heading(f"{clause['number']} — {clause['title']}", level=3)

        meta = doc.add_paragraph()
        r_risk = meta.add_run(f"Risiko: {risk}")
        _set_font(r_risk, bold=True, colour=colour)
        meta.add_run(f"  ·  Kategorie: {clause.get('risk_category', '').upper()}")
        meta.add_run(f"  ·  Seite ~{clause.get('page_start', '?')}")

        if clause.get("reason"):
            doc.add_paragraph(clause["reason"])

        if clause.get("suggestion"):
            sug = doc.add_paragraph()
            r_s = sug.add_run("Empfehlung: ")
            _set_font(r_s, bold=True)
            sug.add_run(clause["suggestion"])

        doc.add_paragraph()

    return _to_bytes(doc)


# ── Mode B — Stellungnahme ────────────────────────────────────────────────────

def export_stellungnahme_docx(data: dict) -> bytes:
    """
    Export Mode B result as a formal Stellungnahme DOCX.

    data keys: nachtrag_summary (dict), positions (list), stellungnahme (str)
    """
    doc = Document()
    _set_page_margins(doc)

    # Title
    title = doc.add_heading("Stellungnahme zum Nachtrag", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph(f"Datum: {datetime.now().strftime('%d.%m.%Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    # ── Financial summary table ──
    summary = data.get("nachtrag_summary", {})
    doc.add_heading("Finanzielle Zusammenfassung", level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    rec = summary.get("recommendation", "negotiate").upper()
    rec_colour = _ASSESS_COLOUR.get(rec, _DARK)

    rows_data = [
        ("Gesamtforderung (AN)", f"EUR {_fmt(summary.get('total_claimed', 0))}"),
        ("Anerkannte Summe", f"EUR {_fmt(summary.get('accepted_total', 0))}"),
        ("Bestrittene Summe", f"EUR {_fmt(summary.get('contested_total', 0))}"),
        ("Anzahl Positionen", str(summary.get("position_count", 0))),
        ("Gesamtempfehlung", rec),
    ]

    for i, (label, value) in enumerate(rows_data):
        label_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        label_cell.text = label

        if i == 4:  # recommendation row
            p = value_cell.paragraphs[0]
            r = p.add_run(value)
            _set_font(r, bold=True, colour=rec_colour)
        else:
            value_cell.text = value

    doc.add_paragraph()

    # ── Stellungnahme text ──
    doc.add_heading("Stellungnahme", level=2)
    stellungnahme = data.get("stellungnahme", "")
    for para_text in stellungnahme.split("\n\n"):
        para_text = para_text.strip()
        if para_text:
            doc.add_paragraph(para_text)

    doc.add_paragraph()
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("Unterschrift / Stempel Auftraggeber")
    doc.add_paragraph()

    # ── Position details ──
    doc.add_page_break()
    doc.add_heading("Positions-Bewertung", level=2)

    for i, pos in enumerate(data.get("positions", []), 1):
        oz = pos.get("oz") or pos.get("lv_oz") or f"Pos. {i}"
        desc = pos.get("nachtrag_description", "")[:70]
        doc.add_heading(f"OZ {oz} — {desc}", level=3)

        assessment = pos.get("assessment", "negotiate").upper()
        a_colour = _ASSESS_COLOUR.get(assessment, _DARK)

        meta = doc.add_paragraph()
        r_a = meta.add_run(f"Bewertung: {assessment}")
        _set_font(r_a, bold=True, colour=a_colour)
        meta.add_run(f"  ·  {pos.get('vob_paragraph', '')}")
        meta.add_run(f"  ·  Forderung: EUR {_fmt(pos.get('nachtrag_claimed_total', 0))}")

        if pos.get("reason"):
            doc.add_paragraph(pos["reason"])

        if pos.get("negotiation_position"):
            neg = doc.add_paragraph()
            r_n = neg.add_run("Verhandlungsposition: ")
            _set_font(r_n, bold=True)
            neg.add_run(pos["negotiation_position"])

        doc.add_paragraph()

    return _to_bytes(doc)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _fmt(value) -> str:
    """Format float as German currency string."""
    try:
        return f"{float(value):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except (TypeError, ValueError):
        return "0,00"


def _set_page_margins(doc: Document, margin_cm: float = 2.5):
    """Set uniform page margins."""
    for section in doc.sections:
        margin = Cm(margin_cm)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def _to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()